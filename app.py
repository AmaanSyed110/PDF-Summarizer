import streamlit as st
import pdfplumber
from openai import OpenAI
import numpy as np
from pathlib import Path
import tempfile
import os
from datetime import datetime
from dotenv import load_dotenv
import faiss
from fpdf import FPDF
from docx import Document

# Load environment variables
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")
if not API_KEY:
    raise ValueError("OPENAI_API_KEY not found in .env file.")

# Initialize session state
if 'current_summaries' not in st.session_state:
    st.session_state.current_summaries = {}
if 'display_output' not in st.session_state:
    st.session_state.display_output = True
if 'history' not in st.session_state:
    st.session_state.history = []
if 'saved_summaries' not in st.session_state:
    st.session_state.saved_summaries = {}
if 'current_view' not in st.session_state:
    st.session_state.current_view = 'main'
if 'settings' not in st.session_state:
    st.session_state.settings = {
        'max_summary_length_words': 200, 
        'chunk_size': 1000
    }
if 'faiss_index' not in st.session_state:
    st.session_state.faiss_index = None
if 'sentence_cache' not in st.session_state:
    st.session_state.sentence_cache = {}

# Initialize OpenAI client
client = OpenAI(api_key=API_KEY)

# Initialize FAISS index globally
dimension = 3072  # Dimension for text-embedding-3-large
if st.session_state.faiss_index is None:
    st.session_state.faiss_index = faiss.IndexFlatL2(dimension)

# Helper functions
def words_to_tokens(words):
    """Convert word count to approximate token count (1 word ≈ 1.3 tokens)."""
    return int(words * 1.3)

def switch_to_history():
    st.session_state.current_view = 'history'

def switch_to_main():
    st.session_state.current_view = 'main'

def extract_text_from_pdf(uploaded_file):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        with pdfplumber.open(tmp_file_path) as pdf:
            text = "".join(page.extract_text() or "" for page in pdf.pages)
        os.unlink(tmp_file_path)
        return text.strip()
    except Exception as e:
        st.error(f"Error extracting text from PDF: {str(e)}")
        return None

def split_text_into_chunks(text, max_chunk_length):
    sentences = text.split('. ')
    chunks = []
    current_chunk = []
    current_length = 0
    for sentence in sentences:
        sentence_length = len(sentence.split())
        if current_length + sentence_length > max_chunk_length:
            chunks.append('. '.join(current_chunk) + '.')
            current_chunk = [sentence]
            current_length = sentence_length
        else:
            current_chunk.append(sentence)
            current_length += sentence_length
    if current_chunk:
        chunks.append('. '.join(current_chunk) + '.')
    return chunks

def get_embedding(text):
    try:
        response = client.embeddings.create(
            model="text-embedding-3-large",
            input=text
        )
        return response.data[0].embedding
    except Exception as e:
        st.error(f"Error getting embedding: {str(e)}")
        return None

def search_relevant_sections(text, query):
    try:
        sentences = text.split('. ')
        query_embedding = get_embedding(query)
        if not query_embedding:
            return text

        # Check cache for existing sentence embeddings
        text_hash = hash(text)
        if text_hash not in st.session_state.sentence_cache:
            sentence_embeddings = []
            valid_sentences = []
            for sentence in sentences:
                if sentence.strip():
                    embedding = get_embedding(sentence)
                    if embedding:
                        sentence_embeddings.append(embedding)
                        valid_sentences.append(sentence)
            if not sentence_embeddings:
                return text
            st.session_state.sentence_cache[text_hash] = {
                'embeddings': np.array(sentence_embeddings).astype('float32'),
                'sentences': valid_sentences
            }
        else:
            sentence_embeddings = st.session_state.sentence_cache[text_hash]['embeddings']
            valid_sentences = st.session_state.sentence_cache[text_hash]['sentences']

        # Reset FAISS index
        st.session_state.faiss_index.reset()
        st.session_state.faiss_index.add(sentence_embeddings)

        # Search for top 5 similar sentences
        query_embedding = np.array([query_embedding]).astype('float32')
        distances, indices = st.session_state.faiss_index.search(query_embedding, k=5)
        relevant_text = '. '.join([valid_sentences[idx] for idx in indices[0]])
        return relevant_text
    except Exception as e:
        st.error(f"Error searching relevant sections: {str(e)}")
        return text

def generate_summary_with_gpt4_chunks(text, max_length_words):
    try:
        max_length_tokens = words_to_tokens(max_length_words)
        chunks = split_text_into_chunks(text, st.session_state.settings['chunk_size'])
        summaries = []
        chunk_max_length_tokens = max_length_tokens // max(1, len(chunks))  # Distribute tokens across chunks

        for chunk in chunks:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that summarizes text concisely."},
                    {"role": "user", "content": f"Summarize this text: {chunk}"}
                ],
                max_tokens=chunk_max_length_tokens,
                temperature=0.7
            )
            summary = response.choices[0].message.content.strip()
            summaries.append(summary)

        # Combine chunk summaries
        combined_summary = " ".join(summaries)
        if len(summaries) > 1:
            # Generate a final concise summary of combined summaries
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that creates a concise summary from multiple summaries."},
                    {"role": "user", "content": f"Create a concise summary from these summaries: {combined_summary}"}
                ],
                max_tokens=max_length_tokens,
                temperature=0.7
            )
            combined_summary = response.choices[0].message.content.strip()

        return combined_summary
    except Exception as e:
        st.error(f"Error generating summary: {str(e)}")
        return None

def save_summary_history(filename, summary):
    st.session_state.history.append({
        'filename': filename,
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'summary': summary
    })

def export_summaries_as_text():
    if st.session_state.history:
        export_text = "PDF SUMMARIES EXPORT\n"
        export_text += f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        export_text += "=" * 50 + "\n\n"
        for entry in st.session_state.history:
            export_text += f"File: {entry['filename']}\n"
            export_text += f"Date: {entry['timestamp']}\n"
            export_text += "-" * 30 + "\n"
            export_text += f"{entry['summary']}\n"
            export_text += "=" * 50 + "\n\n"
        return export_text
    return None

def show_settings_interface():
    with st.sidebar:
        st.header("⚙️ Settings")
        max_summary_length_words = st.slider(
            "Maximum Summary Length (words)",
            min_value=50,
            max_value=500,
            value=st.session_state.settings['max_summary_length_words'],
            step=25
        )
        chunk_size = st.slider(
            "Chunk Size (words)",
            min_value=500,
            max_value=2000,
            value=st.session_state.settings['chunk_size'],
            step=100
        )
        if st.button("Apply Settings"):
            st.session_state.settings['max_summary_length_words'] = max_summary_length_words
            st.session_state.settings['chunk_size'] = chunk_size
            st.success("Settings updated!")
        st.divider()

def show_main_interface():
    container = st.container()
    with container:
        uploaded_files = st.file_uploader("Upload PDF files", type="pdf", accept_multiple_files=True)
        search_query = st.text_input("🔍 Enter keywords to focus on specific topics (optional)")
        if uploaded_files:
            for uploaded_file in uploaded_files:
                st.write(f"### Processing: {uploaded_file.name}")
                file_key = f"{uploaded_file.name}_{search_query}"
                if file_key not in st.session_state.current_summaries:
                    with st.spinner("Extracting text from PDF..."):
                        text = extract_text_from_pdf(uploaded_file)
                    if text:
                        if search_query:
                            with st.spinner("Searching relevant sections..."):
                                text = search_relevant_sections(text, search_query)
                        with st.spinner("Generating summary..."):
                            summary = generate_summary_with_gpt4_chunks(
                                text,
                                st.session_state.settings['max_summary_length_words']
                            )
                        if summary:
                            st.session_state.current_summaries[file_key] = summary
                            save_summary_history(uploaded_file.name, summary)
                    else:
                        st.error(f"Failed to process {uploaded_file.name}")
                        continue
                summary = st.session_state.current_summaries[file_key]
                st.success("Summary generated successfully!")
                st.write(summary)

                col1, col2 = st.columns([1, 2])
                with col1:
                    st.write("Download Summary:")
                    pdf_summary_filename = f"summary_{uploaded_file.name.split('.')[0]}.pdf"
                    try:
                        class UTF8PDF(FPDF):
                            def __init__(self):
                                super().__init__()
                                self.add_font('DejaVu', '', 'DejaVuSansCondensed.ttf', uni=True)
                            def header(self):
                                pass
                            def footer(self):
                                pass
                        pdf = UTF8PDF()
                        pdf.add_page()
                        pdf.set_font('DejaVu', size=12)
                        pdf.multi_cell(0, 10, summary)
                        pdf_output = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                        pdf.output(pdf_output.name)
                    except Exception:
                        pdf = FPDF()
                        pdf.add_page()
                        pdf.set_font('Arial', size=12)
                        summary_ascii = ''.join(char if ord(char) < 128 else '?' for char in summary)
                        pdf.multi_cell(0, 10, summary_ascii)
                        pdf_output = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                        pdf.output(pdf_output.name)
                    st.download_button(
                        label="Download as PDF",
                        data=open(pdf_output.name, "rb").read(),
                        file_name=pdf_summary_filename,
                        mime="application/pdf",
                        key=f"pdf_{file_key}"
                    )
                    docx_summary_filename = f"summary_{uploaded_file.name.split('.')[0]}.docx"
                    doc = Document()
                    doc.add_heading("Summary", level=1)
                    doc.add_paragraph(summary)
                    docx_output = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    doc.save(docx_output.name)
                    st.download_button(
                        label="Download as DOCX",
                        data=open(docx_output.name, "rb").read(),
                        file_name=docx_summary_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"docx_{file_key}"
                    )
                    os.unlink(pdf_output.name)
                    os.unlink(docx_output.name)

def show_history_interface():
    st.header("Summary History")
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("← Back to Main"):
            switch_to_main()
            st.rerun()
    with col2:
        if st.session_state.history:
            export_text = export_summaries_as_text()
            if export_text:
                st.download_button(
                    label="Export All Summaries",
                    data=export_text,
                    file_name="pdf_summaries_export.txt",
                    mime="text/plain"
                )
    if st.session_state.history:
        for entry in st.session_state.history:
            with st.expander(f"📄 {entry['filename']} - {entry['timestamp']}", expanded=True):
                st.write(entry['summary'])
    else:
        st.info("No summaries in history yet. Process some PDFs to see them here!")

def main():
    st.set_page_config(
        page_title="Multi-PDF Summarizer",
        layout="centered",
        initial_sidebar_state="auto"
    )
    show_settings_interface()
    main_container = st.container()
    with main_container:
        st.title("📄 Multi-PDF Summarizer")
        st.subheader("Upload multiple PDFs to generate concise summaries.", divider="gray")
        if st.session_state.current_view == 'main':
            if st.button("📚 View Summary History", type="secondary"):
                switch_to_history()
                st.rerun()
            show_main_interface()
        else:
            show_history_interface()

if __name__ == "__main__":
    main()